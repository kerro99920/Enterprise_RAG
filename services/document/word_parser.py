"""
========================================
Word æ–‡æ¡£è§£æå™¨
========================================

ğŸ“š æ¨¡å—è¯´æ˜ï¼š
- è§£æWordæ–‡æ¡£ï¼ˆ.docx, .docï¼‰
- æå–æ–‡æœ¬ã€è¡¨æ ¼ã€å›¾ç‰‡ç­‰å…ƒç´ 
- ä¿ç•™æ–‡æ¡£ç»“æ„ä¿¡æ¯

ğŸ¯ æ ¸å¿ƒåŠŸèƒ½ï¼š
1. æ–‡æœ¬æå–ï¼ˆä¿ç•™æ®µè½ç»“æ„ï¼‰
2. è¡¨æ ¼è¯†åˆ«å’Œæå–
3. æ ·å¼ä¿¡æ¯æå–
4. å…ƒæ•°æ®æå–

========================================
"""
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.logger import logger, log_execution
from core.config import settings


class WordParser:
    """
    Wordæ–‡æ¡£è§£æå™¨

    ğŸ¯ èŒè´£ï¼š
    - æå–Wordæ–‡æ¡£æ–‡æœ¬
    - æå–è¡¨æ ¼æ•°æ®
    - æå–æ–‡æ¡£å…ƒæ•°æ®
    - ä¿ç•™æ–‡æ¡£ç»“æ„
    """

    def __init__(self):
        """åˆå§‹åŒ–Wordè§£æå™¨"""
        self.supported_extensions = ['.docx', '.doc']

    @log_execution("è§£æWordæ–‡æ¡£")
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        è§£æWordæ–‡ä»¶

        å‚æ•°ï¼š
            file_path: Wordæ–‡ä»¶è·¯å¾„

        è¿”å›ï¼š
            Dict: è§£æç»“æœ
            {
                "text": str,              # å…¨æ–‡æœ¬
                "paragraphs": List[Dict], # æ®µè½åˆ—è¡¨
                "tables": List[Dict],     # è¡¨æ ¼æ•°æ®
                "metadata": Dict,         # æ–‡æ¡£å…ƒæ•°æ®
                "has_tables": bool,       # æ˜¯å¦åŒ…å«è¡¨æ ¼
                "paragraph_count": int,   # æ®µè½æ•°
                "word_count": int,        # å­—æ•°
            }
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")

            logger.info(f"å¼€å§‹è§£æWord: {file_path}")

            # æ‰“å¼€Wordæ–‡æ¡£
            doc = Document(file_path)

            # æå–å…ƒæ•°æ®
            metadata = self._extract_metadata(doc)

            # æå–æ®µè½
            paragraphs_data = self._extract_paragraphs(doc)

            # æå–è¡¨æ ¼
            tables = self._extract_tables(doc)

            # åˆå¹¶æ‰€æœ‰æ–‡æœ¬
            all_text = "\n\n".join([p["text"] for p in paragraphs_data if p["text"]])

            # ç»Ÿè®¡å­—æ•°
            word_count = len(all_text.replace(" ", "").replace("\n", ""))

            result = {
                "text": all_text,
                "paragraphs": paragraphs_data,
                "tables": tables,
                "metadata": metadata,
                "has_tables": len(tables) > 0,
                "paragraph_count": len(paragraphs_data),
                "word_count": word_count,
                "file_path": file_path,
                "file_name": os.path.basename(file_path)
            }

            logger.info(
                f"Wordè§£æå®Œæˆ: {file_path} | "
                f"æ®µè½: {len(paragraphs_data)} | "
                f"è¡¨æ ¼: {len(tables)} | "
                f"å­—æ•°: {word_count}"
            )

            return result

        except Exception as e:
            logger.error(f"Wordè§£æå¤±è´¥: {file_path} | é”™è¯¯: {str(e)}")
            raise

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """
        æå–Wordæ–‡æ¡£å…ƒæ•°æ®

        å‚æ•°ï¼š
            doc: Documentå¯¹è±¡

        è¿”å›ï¼š
            Dict: å…ƒæ•°æ®ä¿¡æ¯
        """
        try:
            # Wordæ ¸å¿ƒå±æ€§
            core_props = doc.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
            }

            return metadata

        except Exception as e:
            logger.warning(f"æå–Wordå…ƒæ•°æ®å¤±è´¥: {str(e)}")
            return {}

    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        æå–æ®µè½ä¿¡æ¯

        å‚æ•°ï¼š
            doc: Documentå¯¹è±¡

        è¿”å›ï¼š
            List[Dict]: æ®µè½æ•°æ®åˆ—è¡¨

        ğŸ’¡ æ®µè½æ ¼å¼ï¼š
        {
            "index": int,         # æ®µè½ç´¢å¼•
            "text": str,          # æ®µè½æ–‡æœ¬
            "style": str,         # æ ·å¼åç§°
            "is_heading": bool,   # æ˜¯å¦ä¸ºæ ‡é¢˜
            "level": int,         # æ ‡é¢˜çº§åˆ«ï¼ˆå¦‚æœæ˜¯æ ‡é¢˜ï¼‰
        }
        """
        paragraphs_data = []

        try:
            for index, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                # è·³è¿‡ç©ºæ®µè½
                if not text:
                    continue

                # è·å–æ ·å¼ä¿¡æ¯
                style_name = paragraph.style.name if paragraph.style else "Normal"

                # åˆ¤æ–­æ˜¯å¦ä¸ºæ ‡é¢˜
                is_heading = style_name.startswith("Heading")

                # æå–æ ‡é¢˜çº§åˆ«
                level = 0
                if is_heading:
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except:
                        level = 0

                para_data = {
                    "index": index,
                    "text": text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "level": level
                }

                paragraphs_data.append(para_data)

            return paragraphs_data

        except Exception as e:
            logger.error(f"æå–Wordæ®µè½å¤±è´¥: {str(e)}")
            raise

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        æå–Wordæ–‡æ¡£ä¸­çš„è¡¨æ ¼

        å‚æ•°ï¼š
            doc: Documentå¯¹è±¡

        è¿”å›ï¼š
            List[Dict]: è¡¨æ ¼æ•°æ®åˆ—è¡¨

        ğŸ’¡ è¡¨æ ¼æ ¼å¼ï¼š
        {
            "table_index": int,   # è¡¨æ ¼ç´¢å¼•
            "rows": int,          # è¡Œæ•°
            "cols": int,          # åˆ—æ•°
            "data": List[List],   # è¡¨æ ¼æ•°æ®
            "text": str           # è¡¨æ ¼è½¬æ¢ä¸ºæ–‡æœ¬
        }
        """
        tables = []

        try:
            for table_index, table in enumerate(doc.tables):
                # æå–è¡¨æ ¼æ•°æ®
                table_data = []

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)

                # è½¬æ¢ä¸ºæ–‡æœ¬
                table_text = self._table_to_text(table_data)

                table_info = {
                    "table_index": table_index,
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0,
                    "data": table_data,
                    "text": table_text
                }

                tables.append(table_info)

            logger.info(f"æå–åˆ° {len(tables)} ä¸ªè¡¨æ ¼")
            return tables

        except Exception as e:
            logger.warning(f"æå–Wordè¡¨æ ¼å¤±è´¥: {str(e)}")
            return []

    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """
        å°†è¡¨æ ¼æ•°æ®è½¬æ¢ä¸ºæ–‡æœ¬æ ¼å¼

        å‚æ•°ï¼š
            table_data: è¡¨æ ¼æ•°æ®ï¼ˆäºŒç»´åˆ—è¡¨ï¼‰

        è¿”å›ï¼š
            str: æ ¼å¼åŒ–çš„è¡¨æ ¼æ–‡æœ¬
        """
        if not table_data:
            return ""

        try:
            lines = []
            for row in table_data:
                # ç”¨ | åˆ†éš”å•å…ƒæ ¼
                line = " | ".join(row)
                lines.append(line)

            return "\n".join(lines)

        except Exception as e:
            logger.warning(f"è¡¨æ ¼è½¬æ–‡æœ¬å¤±è´¥: {str(e)}")
            return ""

    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        æå–æ–‡æ¡£ä¸­çš„æ‰€æœ‰æ ‡é¢˜

        å‚æ•°ï¼š
            file_path: Wordæ–‡ä»¶è·¯å¾„

        è¿”å›ï¼š
            List[Dict]: æ ‡é¢˜åˆ—è¡¨

        ğŸ’¡ ç”¨é€”ï¼š
        - ç”Ÿæˆæ–‡æ¡£ç›®å½•
        - ç†è§£æ–‡æ¡£ç»“æ„
        - è¾…åŠ©æ–‡æœ¬åˆ†å—
        """
        try:
            doc = Document(file_path)
            headings = []

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    try:
                        level = int(paragraph.style.name.replace("Heading", "").strip())
                    except:
                        level = 0

                    headings.append({
                        "text": paragraph.text.strip(),
                        "level": level,
                        "style": paragraph.style.name
                    })

            return headings

        except Exception as e:
            logger.error(f"æå–æ ‡é¢˜å¤±è´¥: {str(e)}")
            return []

    def get_text_with_structure(self, file_path: str) -> str:
        """
        è·å–ä¿ç•™ç»“æ„çš„æ–‡æœ¬

        å‚æ•°ï¼š
            file_path: Wordæ–‡ä»¶è·¯å¾„

        è¿”å›ï¼š
            str: å¸¦ç»“æ„æ ‡è®°çš„æ–‡æœ¬

        ğŸ’¡ æ ¼å¼ç¤ºä¾‹ï¼š
        # æ ‡é¢˜1
        æ­£æ–‡æ®µè½...

        ## æ ‡é¢˜2
        æ­£æ–‡æ®µè½...

        [è¡¨æ ¼]
        | åˆ—1 | åˆ—2 |
        | æ•°æ®1 | æ•°æ®2 |
        """
        try:
            result = self.parse(file_path)

            structured_text = []

            # æ·»åŠ æ®µè½ï¼ˆæ ‡é¢˜ç”¨#æ ‡è®°ï¼‰
            for para in result['paragraphs']:
                if para['is_heading']:
                    # æ ‡é¢˜å‰åŠ #å·
                    prefix = "#" * para['level']
                    structured_text.append(f"{prefix} {para['text']}")
                else:
                    structured_text.append(para['text'])

            # æ·»åŠ è¡¨æ ¼
            for table in result['tables']:
                structured_text.append("\n[è¡¨æ ¼]")
                structured_text.append(table['text'])

            return "\n\n".join(structured_text)

        except Exception as e:
            logger.error(f"è·å–ç»“æ„åŒ–æ–‡æœ¬å¤±è´¥: {str(e)}")
            raise


# =========================================
# ğŸ’¡ ä½¿ç”¨ç¤ºä¾‹
# =========================================
"""
# 1. åŸºç¡€ä½¿ç”¨
from services.document.word_parser import WordParser

parser = WordParser()

# è§£æWordæ–‡æ¡£
result = parser.parse("data/raw_docs/é¡¹ç›®æ€»ç»“.docx")

print(f"æ–‡æ¡£æ ‡é¢˜: {result['metadata']['title']}")
print(f"ä½œè€…: {result['metadata']['author']}")
print(f"æ®µè½æ•°: {result['paragraph_count']}")
print(f"å­—æ•°: {result['word_count']}")
print(f"è¡¨æ ¼æ•°: {len(result['tables'])}")

# è®¿é—®å…¨æ–‡æœ¬
full_text = result['text']
print(full_text[:500])

# è®¿é—®æ®µè½
for para in result['paragraphs'][:5]:
    if para['is_heading']:
        print(f"[æ ‡é¢˜{para['level']}] {para['text']}")
    else:
        print(para['text'][:100])

# è®¿é—®è¡¨æ ¼
for table in result['tables']:
    print(f"\nè¡¨æ ¼ {table['table_index']}:")
    print(table['text'])


# 2. æå–æ ‡é¢˜ï¼ˆç”Ÿæˆç›®å½•ï¼‰
headings = parser.extract_headings("document.docx")
for heading in headings:
    indent = "  " * (heading['level'] - 1)
    print(f"{indent}- {heading['text']}")


# 3. è·å–ç»“æ„åŒ–æ–‡æœ¬
structured_text = parser.get_text_with_structure("document.docx")
print(structured_text)
"""